#!/usr/bin/env python3
"""
doc_read_tool.py - Tool for reading Word document content

This tool can read the content of Microsoft Word documents (.docx), including text and basic formatting information.
"""

import os
from typing import Optional
from langchain.tools import BaseTool
from pydantic import BaseModel, Field

try:
    from docx import Document
except ImportError:
    print("Error: python-docx package is not installed. Please run: pip install python-docx")
    exit(1)


class DocReadToolInput(BaseModel):
    file_path: str = Field(..., description="Absolute path to the Word document (.docx)")
    max_length: Optional[int] = Field(None, description="Optional, limit the length of returned text")


class DocReadTool(BaseTool):
    name: str = "doc_read"
    description: str = (
        "Read the content of Microsoft Word documents (.docx), including text and table data. "
        "Can extract all paragraph text and table content from the document. "
        "Supports content length limitation to avoid processing overly large documents."
    )
    args_schema: type = DocReadToolInput

    def _run(self, file_path: str, max_length: Optional[int] = None) -> str:
        """
        Read the content of a Word document (.docx)
        
        Args:
            file_path: Path to the Word document
            max_length: Optional, limit the length of returned text
            
        Returns:
            Text content of the document
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return f"Error: File '{file_path}' does not exist"
            
            # Check file extension
            if not file_path.lower().endswith('.docx'):
                return f"Error: File '{file_path}' is not a valid .docx file"
            
            # Open document
            doc = Document(file_path)
            
            # Extract text content
            full_text = []
            
            # Read paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Read table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Merge all text
            content = "\n\n".join(full_text)
            
            # If max length is specified, truncate content
            if max_length and len(content) > max_length:
                content = content[:max_length] + "\n\n[Content truncated...]"
            
            return content
        
        except Exception as e:
            return f"Error reading document: {str(e)}"


# For backward compatibility, keep the original function
def read_docx(file_path: str, max_length: Optional[int] = None) -> str:
    """Backward compatible function interface"""
    tool = DocReadTool()
    return tool._run(file_path, max_length)


if __name__ == "__main__":
    import argparse
    
    # Command line entry point
    parser = argparse.ArgumentParser(description="Read Word document content")
    parser.add_argument("file_path", help="Path to the Word document (.docx)")
    parser.add_argument("--max-length", type=int, help="Limit the length of returned text")
    
    args = parser.parse_args()
    
    content = read_docx(args.file_path, args.max_length)
    print(content)